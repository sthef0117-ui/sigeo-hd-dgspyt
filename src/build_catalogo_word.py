"""
SIGEO-HD DGSPYT — Catálogo territorial en Word editable.

Escribe el texto del mapa mural: la tabla Coordinación Regional / Subdirección /
Municipio, completa y sin cifras operativas. El mapa no va como imagen; va su
contenido, para poder editarlo.

    python src/etl_sigeo.py
    python src/build_catalogo_word.py

Salida: entregables/CATALOGO_TERRITORIAL_DGSPYT.docx

Documento editable de verdad: estilos normales de Word, una tabla real, sin
imágenes ni cajas de texto flotantes, para que el área pueda corregir un
municipio mal ubicado sin pelearse con el formato.

Cobertura: se listan todos los municipios presentes en la cartografía del
estado y en el inventario de instalaciones, no solo los que tuvieron
incidencia en el corte.
"""

import json
import sys
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RAIZ = Path(__file__).resolve().parent.parent
ANALISIS = RAIZ / "analisis"
POLIGONOS = RAIZ / "poligonos" / "municipios_edomex.geojson"
SALIDA = RAIZ / "entregables"

GUINDA = RGBColor(0x7A, 0x13, 0x27)
TINTA = RGBColor(0x16, 0x20, 0x2C)
GRIS = RGBColor(0x59, 0x65, 0x73)

MENORES = {"de", "del", "la", "las", "el", "los", "en", "y", "con", "por", "a", "al"}

# El pipeline normaliza el nombre para poder cruzar las fuentes: quita acentos y
# recorta sufijos. Para un catalogo eso no sirve, asi que aqui se restituye el
# nombre oficial. Los que la cartografia no trae van en esta tabla.
NOMBRE_OFICIAL = {
    "ATIZAPAN": "Atizapán de Zaragoza",
    "ECATEPEC": "Ecatepec de Morelos",
    "NAUCALPAN": "Naucalpan de Juárez",
    "TLALNEPANTLA": "Tlalnepantla de Baz",
    "VALLE DE CHALCO": "Valle de Chalco Solidaridad",
    "TONANITLA": "Tonanitla",
    "SAN JOSE DEL RINCON": "San José del Rincón",
    "NEZAHUALCOYOTL": "Nezahualcóyotl",
    "CUAUTITLAN": "Cuautitlán",
    "CUAUTITLAN IZCALLI": "Cuautitlán Izcalli",
}


def titulo_es(t):
    palabras = str(t).lower().split()
    return " ".join(p.capitalize() if i == 0 or p not in MENORES else p
                    for i, p in enumerate(palabras))


def nombres_cartografia():
    """Nombre oficial con acentos, tal como viene en la cartografía."""
    if not POLIGONOS.exists():
        return {}
    geo = json.loads(POLIGONOS.read_text(encoding="utf-8"))
    return {f["properties"]["municipio"]: f["properties"]["nombre"]
            for f in geo["features"] if f["properties"].get("nombre")}


def leer(nombre):
    ruta = ANALISIS / nombre
    if not ruta.exists():
        print(f"Falta {ruta.name}. Ejecuta primero: python src/etl_sigeo.py")
        sys.exit(1)
    return json.loads(ruta.read_text(encoding="utf-8"))


def sombrear(celda, hex_color):
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:val"), "clear")
    sombra.set(qn("w:fill"), hex_color)
    celda._tc.get_or_add_tcPr().append(sombra)


def escribir(celda, texto, tam=9.5, negrita=False, color=TINTA,
             alineacion=WD_ALIGN_PARAGRAPH.LEFT):
    celda.text = ""
    p = celda.paragraphs[0]
    p.alignment = alineacion
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(str(texto))
    r.font.size = Pt(tam)
    r.font.bold = negrita
    r.font.color.rgb = color
    r.font.name = "Calibri"


def parrafo(doc, texto, tam=10, negrita=False, color=TINTA, espacio=6,
            cursiva=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(espacio)
    r = p.add_run(texto)
    r.font.size = Pt(tam)
    r.font.bold = negrita
    r.font.italic = cursiva
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def completar_por_vecindad(asignado, geo):
    """
    Extiende la coordinación a los municipios que la cartografía tiene y el
    inventario no menciona, por colindancia de polígonos.
    """
    vertices = defaultdict(list)
    for f in geo["features"]:
        nombre = f["properties"]["municipio"]
        g = f["geometry"]
        anillos = (g["coordinates"] if g["type"] == "Polygon"
                   else [a for p in g["coordinates"] for a in p])
        for anillo in anillos:
            vertices[nombre].extend((p[1], p[0]) for p in anillo)

    PASO = 0.02
    indice = defaultdict(set)
    for nombre, puntos in vertices.items():
        for lat, lng in puntos:
            indice[(round(lat / PASO), round(lng / PASO))].add(nombre)

    inferidos = set()
    for _ in range(5):
        pendientes = [m for m in vertices if m not in asignado]
        if not pendientes:
            break
        avance = False
        for m in pendientes:
            votos = Counter()
            for lat, lng in vertices[m]:
                celda = (round(lat / PASO), round(lng / PASO))
                for dy in (-1, 0, 1):
                    for dx in (-1, 0, 1):
                        for vecino in indice.get((celda[0] + dy, celda[1] + dx), ()):
                            if vecino != m and vecino in asignado:
                                votos[asignado[vecino][0]] += 1
            if votos:
                asignado[m] = (votos.most_common(1)[0][0], "", "vecindad")
                inferidos.add(m)
                avance = True
        if not avance:
            break
    return inferidos


def main():
    territorio = leer("perfil_territorial.json")
    coord = leer("perfil_coordinaciones.json")["coordinaciones"]

    # municipio -> (coordinacion, subdireccion, origen)
    asignado = {t["municipio"]: (t["coordinacion"], t["subdireccion"],
                                 t["coordinacion_origen"])
                for t in territorio if t["coordinacion"]}

    inferidos_extra = set()
    if POLIGONOS.exists():
        geo = json.loads(POLIGONOS.read_text(encoding="utf-8"))
        inferidos_extra = completar_por_vecindad(asignado, geo)
    oficiales = nombres_cartografia()

    def nombre_oficial(clave):
        return (NOMBRE_OFICIAL.get(clave)
                or oficiales.get(clave)
                or titulo_es(clave))

    filas = []
    for municipio, (c, s, origen) in asignado.items():
        filas.append({
            "coordinacion": c,
            "subdireccion": s or "Sin subdirección declarada",
            "municipio": nombre_oficial(municipio),
            "origen": origen,
        })
    orden_coord = {c["coordinacion"]: i for i, c in enumerate(coord)}
    filas.sort(key=lambda f: (orden_coord.get(f["coordinacion"], 99),
                              f["subdireccion"].startswith("Sin"),
                              f["subdireccion"], f["municipio"]))

    # ------------------------------------------------------------------
    doc = Document()
    for seccion in doc.sections:
        seccion.top_margin = seccion.bottom_margin = Cm(1.8)
        seccion.left_margin = seccion.right_margin = Cm(2)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    parrafo(doc, "DIRECCIÓN GENERAL DE SEGURIDAD PÚBLICA Y TRÁNSITO",
            tam=10.5, negrita=True, color=GUINDA, espacio=2)
    t = doc.add_heading("Catálogo territorial · Coordinaciones regionales", level=0)
    for r in t.runs:
        r.font.color.rgb = TINTA
    parrafo(doc,
            f"{len(filas)} municipios · {len(coord)} coordinaciones regionales · "
            f"{len({f['subdireccion'] for f in filas if not f['subdireccion'].startswith('Sin')})} "
            "subdirecciones",
            tam=9.5, color=GRIS, espacio=10, cursiva=True)

    usa_v = any(f["origen"] == "vecindad" for f in filas)
    usa_e = any(f["origen"] == "escision" for f in filas)
    glosa = ["Estructura tomada del inventario oficial de instalaciones de la Dirección."]
    if usa_v:
        glosa.append("Los municipios marcados con (v) no tienen instalación propia "
                     "y su coordinación se asignó por colindancia territorial.")
    if usa_e:
        glosa.append("Los marcados con (e) la heredan del municipio del que se "
                     "escindieron.")
    if usa_v or usa_e:
        glosa.append("Son inferencias y conviene confirmarlas.")
    parrafo(doc, " ".join(glosa), tam=9.5, color=GRIS, espacio=12)

    # --- Tabla única, a dos columnas de página para que quepa ---
    seccion = doc.sections[-1]
    cols = seccion._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "340")

    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    anchos = (Cm(2.9), Cm(2.6), Cm(3.3))
    for j, encabezado in enumerate(["Coordinación", "Subdirección", "Municipio"]):
        escribir(tabla.rows[0].cells[j], encabezado, tam=9, negrita=True,
                 color=RGBColor(0xFF, 0xFF, 0xFF))
        sombrear(tabla.rows[0].cells[j], "16202C")
        tabla.columns[j].width = anchos[j]

    marca = {"vecindad": " (v)", "escision": " (e)"}
    coord_previa = sub_previa = None
    for f in filas:
        celdas = tabla.add_row().cells
        nueva_coord = f["coordinacion"] != coord_previa
        nueva_sub = nueva_coord or f["subdireccion"] != sub_previa
        escribir(celdas[0], titulo_es(f["coordinacion"]) if nueva_coord else "",
                 tam=9, negrita=True, color=GUINDA)
        escribir(celdas[1], titulo_es(f["subdireccion"]) if nueva_sub else "",
                 tam=9, color=GRIS)
        escribir(celdas[2], f["municipio"] + marca.get(f["origen"], ""), tam=9)
        for j in range(3):
            celdas[j].width = anchos[j]
        if nueva_coord:
            for j in range(3):
                sombrear(celdas[j], "F2F4F7")
        coord_previa, sub_previa = f["coordinacion"], f["subdireccion"]

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.sections[-1]._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "1")

    parrafo(doc, "", espacio=4)
    parrafo(doc,
            "Los números impresos sobre los polígonos del mapa mural son las claves "
            "municipales del INEGI. No se incluyen en esta tabla porque el catálogo "
            "oficial de claves no está entre los insumos entregados; si se "
            "proporciona, se agregan como columna.",
            tam=9, color=GRIS, espacio=4, cursiva=True)

    SALIDA.mkdir(parents=True, exist_ok=True)
    destino = SALIDA / "CATALOGO_TERRITORIAL_DGSPYT.docx"
    doc.save(destino)

    inferidos = sum(1 for f in filas if f["origen"] != "catalogo")
    print("SIGEO-HD DGSPYT · catálogo territorial en Word")
    print(f"  municipios en la tabla ..... {len(filas)}")
    print(f"  del inventario oficial ..... {len(filas) - inferidos}")
    print(f"  inferidos .................. {inferidos} "
          f"({len(inferidos_extra)} añadidos desde la cartografía)")
    print(f"  archivo .................... {destino}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
