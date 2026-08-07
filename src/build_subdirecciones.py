"""
SIGEO-HD DGSPYT — Listado de subdirecciones con ubicación y coordinación.

Lee el inventario de instalaciones y extrae únicamente las subdirecciones:
cuántas son, cuál es la ubicación de su comandancia sede y a qué coordinación
regional pertenece cada una.

    python src/build_subdirecciones.py

Salidas en entregables/:
    SUBDIRECCIONES_DGSPYT.docx    listado en Word editable
    SUBDIRECCIONES_DGSPYT.csv     el mismo listado en tabla

La sede se identifica por el inmueble cuya unidad usuaria es la comandancia de
esa subdireccion. Cuando el inventario no declara comandancia propia, se toma
el inmueble de esa subdireccion con mayor personal adscrito y se marca, porque
es una deduccion y no un dato declarado.
"""

import csv
import re
import sys
import unicodedata
from collections import defaultdict
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RAIZ = Path(__file__).resolve().parent.parent
XLS = RAIZ / "insumos" / "estructura" / "INSTALACIONES_DGSPYT_290426_QR L.xlsx"
SALIDA = RAIZ / "entregables"

GUINDA = RGBColor(0x7A, 0x13, 0x27)
TINTA = RGBColor(0x16, 0x20, 0x2C)
GRIS = RGBColor(0x59, 0x65, 0x73)

# Unidades del inventario que no mandan territorio.
NO_TERRITORIALES = {"MONTADOS CANININOS Y GAMA", "MONTADOS CANINOS Y GAMA",
                    "TRANSITO IXTAPAN", "TRANSITO METROPOLITANO"}

# El inventario escribe la misma subdireccion de dos formas.
ALIAS = {"CEM": "CIRCUITO EXTERIOR MEXIQUENSE"}

MENORES = {"de", "del", "la", "las", "el", "los", "en", "y", "con", "por", "a", "al"}


def norm(t):
    s = unicodedata.normalize("NFD", str(t or ""))
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return " ".join(s.upper().split())


def titulo_es(t):
    palabras = str(t).lower().split()
    return " ".join(p.capitalize() if i == 0 or p not in MENORES else p
                    for i, p in enumerate(palabras))


def limpio(v):
    s = re.sub(r"\s+", " ", str(v or "")).strip()
    return "" if s.upper() in ("", "0", "NONE", "N/A", "SIN INFORMACION") else s


def entero(v):
    try:
        return int(float(str(v)))
    except (TypeError, ValueError):
        return 0


def sombrear(celda, color):
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:val"), "clear")
    sombra.set(qn("w:fill"), color)
    celda._tc.get_or_add_tcPr().append(sombra)


def escribir(celda, texto, tam=9, negrita=False, color=TINTA):
    celda.text = ""
    p = celda.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(str(texto))
    r.font.size = Pt(tam)
    r.font.bold = negrita
    r.font.color.rgb = color
    r.font.name = "Calibri"


def parrafo(doc, texto, tam=10, negrita=False, color=TINTA, espacio=6,
            cursiva=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(espacio)
    r = p.add_run(texto)
    r.font.size = Pt(tam)
    r.font.bold = negrita
    r.font.italic = cursiva
    r.font.color.rgb = color
    r.font.name = "Calibri"


def cargar():
    wb = openpyxl.load_workbook(XLS, read_only=True, data_only=True)
    ws = wb[wb.sheetnames[0]]
    inmuebles = []
    for f in ws.iter_rows(min_row=7, values_only=True):
        cvo = limpio(f[0])
        if not cvo or not cvo.replace(".0", "").isdigit():
            continue
        sub = ALIAS.get(norm(f[2]), norm(f[2]))
        inmuebles.append({
            "cvo": int(float(cvo)),
            "coordinacion": norm(f[1]),
            "subdireccion": sub,
            "municipio": limpio(f[3]),
            "ubicacion": limpio(f[4]),
            "unidad": limpio(f[5]),
            "cp": limpio(f[6]),
            "coordenadas": limpio(f[7]),
            "en_uso": limpio(f[12]).upper().startswith("SI"),
            "personal": entero(f[16]),
        })
    wb.close()
    return inmuebles


def es_sede(inmueble, nombre_sub=None):
    """
    ¿Este inmueble aloja la comandancia de su propia subdirección?

    El inventario nombra la misma sede de dos maneras: unas veces
    "COMANDANCIA DE LA SUBDIRECCION OPERATIVA REGIONAL <X>" y otras
    "COMANDANCIA DEL AREA REGIONAL <X>". Se aceptan ambas.
    """
    u = norm(inmueble["unidad"])
    sub = nombre_sub or inmueble["subdireccion"]
    if not sub or "COMANDANCIA" not in u:
        return False
    if sub.split()[0] not in u:
        return False
    return "SUBDIRECCION" in u or "AREA REGIONAL" in u


def main():
    if not XLS.exists():
        print(f"Falta el inventario: {XLS}")
        return 1

    inmuebles = cargar()
    por_sub = defaultdict(list)
    for i in inmuebles:
        s = i["subdireccion"]
        if s and s not in ("N/A",) and s not in NO_TERRITORIALES:
            por_sub[s].append(i)

    filas = []
    for sub, lista in por_sub.items():
        sedes = [i for i in lista if es_sede(i)]
        if not sedes:
            sedes = [i for i in inmuebles if es_sede(i, sub)
                     and i["coordinacion"] in {x["coordinacion"] for x in lista}]
        if sedes:
            sede = max(sedes, key=lambda i: i["personal"])
            origen = "declarada"
        else:
            sede = max(lista, key=lambda i: i["personal"])
            origen = "deducida"
        coords = sorted({i["coordinacion"] for i in lista if i["coordinacion"]})
        filas.append({
            "subdireccion": sub,
            "coordinacion": sede["coordinacion"] or (coords[0] if coords else ""),
            "otras_coordinaciones": ", ".join(c for c in coords
                                              if c != sede["coordinacion"]),
            "municipio_sede": sede["municipio"],
            "ubicacion_sede": sede["ubicacion"],
            "cp": sede["cp"],
            "coordenadas": sede["coordenadas"],
            "inmuebles": len(lista),
            "personal_total": sum(i["personal"] for i in lista),
            "origen_sede": origen,
        })
    filas.sort(key=lambda f: (f["coordinacion"], f["subdireccion"]))

    SALIDA.mkdir(parents=True, exist_ok=True)

    # ---------------- CSV ----------------
    csv_ruta = SALIDA / "SUBDIRECCIONES_DGSPYT.csv"
    campos = ["subdireccion", "coordinacion", "municipio_sede", "ubicacion_sede",
              "cp", "coordenadas", "inmuebles", "personal_total",
              "otras_coordinaciones", "origen_sede"]
    with csv_ruta.open("w", encoding="utf-8-sig", newline="") as fh:
        w = csv.DictWriter(fh, fieldnames=campos, extrasaction="ignore")
        w.writeheader()
        w.writerows(filas)

    # ---------------- Word ----------------
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.8)
        s.left_margin = s.right_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    parrafo(doc, "DIRECCIÓN GENERAL DE SEGURIDAD PÚBLICA Y TRÁNSITO",
            tam=10.5, negrita=True, color=GUINDA, espacio=2)
    t = doc.add_heading("Subdirecciones operativas regionales", level=0)
    for r in t.runs:
        r.font.color.rgb = TINTA
    parrafo(doc,
            f"{len(filas)} subdirecciones · ubicación de su comandancia sede y "
            "coordinación regional a la que pertenecen",
            tam=9.5, color=GRIS, espacio=12, cursiva=True)

    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    encabezados = ["Subdirección", "Coordinación", "Municipio sede",
                   "Ubicación de la comandancia sede", "Inmuebles"]
    anchos = (Cm(3.4), Cm(2.9), Cm(2.9), Cm(7.4), Cm(1.6))
    for j, h in enumerate(encabezados):
        escribir(tabla.rows[0].cells[j], h, tam=9, negrita=True,
                 color=RGBColor(0xFF, 0xFF, 0xFF))
        sombrear(tabla.rows[0].cells[j], "16202C")
        tabla.columns[j].width = anchos[j]

    coord_previa = None
    for f in filas:
        celdas = tabla.add_row().cells
        marca = "" if f["origen_sede"] == "declarada" else " *"
        escribir(celdas[0], titulo_es(f["subdireccion"]) + marca, negrita=True)
        escribir(celdas[1],
                 titulo_es(f["coordinacion"]) if f["coordinacion"] != coord_previa else "",
                 color=GUINDA, negrita=True)
        escribir(celdas[2], titulo_es(f["municipio_sede"]))
        escribir(celdas[3], f["ubicacion_sede"] +
                 (f"  ·  CP {f['cp']}" if f["cp"] else ""), tam=8.5)
        escribir(celdas[4], f["inmuebles"])
        for j in range(5):
            celdas[j].width = anchos[j]
        coord_previa = f["coordinacion"]

    deducidas = [f for f in filas if f["origen_sede"] != "declarada"]
    parrafo(doc, "", espacio=6)
    if deducidas:
        parrafo(doc,
                "* Sede deducida: el inventario no declara comandancia propia para "
                "esa subdirección, así que se tomó su inmueble con mayor personal "
                "adscrito. Conviene confirmarla.",
                tam=9, color=GUINDA, espacio=4)
    parrafo(doc,
            "Fuente: INSTALACIONES_DGSPYT_290426. Se excluyen las unidades no "
            "territoriales (Montados, Caninos y GAMA, y las de Tránsito), que "
            "aparecen en la columna de subdirección del inventario pero no mandan "
            "territorio. «Circuito Exterior Mexiquense» y «CEM» se unificaron por "
            "ser la misma subdirección escrita de dos formas.",
            tam=9, color=GRIS, espacio=4, cursiva=True)

    docx_ruta = SALIDA / "SUBDIRECCIONES_DGSPYT.docx"
    doc.save(docx_ruta)

    print("SIGEO-HD DGSPYT · listado de subdirecciones")
    print(f"  subdirecciones ............. {len(filas)}")
    print(f"  con sede declarada ......... {len(filas) - len(deducidas)}")
    print(f"  con sede deducida .......... {len(deducidas)}")
    print(f"  Word ....................... {docx_ruta}")
    print(f"  CSV ........................ {csv_ruta}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
